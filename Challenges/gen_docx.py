#!/usr/bin/env python3
"""
Generate DOCX report from BaoCao_ParallelComputing.md
Requires: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

INPUT_FILE = "BaoCao_ParallelComputing.md"
OUTPUT_FILE = "BaoCao_ParallelComputing.docx"


def add_horizontal_line(doc):
    """Add a horizontal line (paragraph border) to the document."""
    p = doc.add_paragraph()
    p_fmt = p.paragraph_format
    p_fmt.space_before = Pt(0)
    p_fmt.space_after = Pt(0)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4472C4")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_bg(cell, color):
    """Set cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tcPr.append(shd)


def process_markdown(doc, lines):
    i = 0
    in_table = False
    table = None
    headers = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # End of table
        if in_table and not stripped.startswith("|"):
            in_table = False
            table = None
            headers = []
            continue

        # Skip horizontal rules inside table blocks
        if in_table and stripped.startswith("---"):
            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            add_horizontal_line(doc)
            heading_text = stripped[2:].strip()
            p = doc.add_heading(heading_text, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
        elif stripped.startswith("## "):
            add_horizontal_line(doc)
            heading_text = stripped[3:].strip()
            doc.add_heading(heading_text, level=2)
            i += 1
        elif stripped.startswith("### "):
            heading_text = stripped[4:].strip()
            doc.add_heading(heading_text, level=3)
            i += 1
        elif stripped.startswith("#### "):
            heading_text = stripped[5:].strip()
            doc.add_heading(heading_text, level=4)
            i += 1

        # Table
        elif stripped.startswith("|") and stripped.endswith("|"):
            # Parse table rows
            if not in_table:
                in_table = True
                # First row: header
                cells = [c.strip() for c in stripped.split("|")[1:-1]]
                headers = cells
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = "Table Grid"
                # Header row
                hdr_row = table.rows[0]
                for j, cell_text in enumerate(cells):
                    cell = hdr_row.cells[j]
                    cell.text = cell_text
                    cell.paragraphs[0].runs[0].bold = True
                    set_cell_bg(cell, "D6E4F0")
                i += 1
                # Skip separator row
                if i < len(lines) and "---" in lines[i]:
                    i += 1
                continue

            # Data row
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            if len(cells) == len(headers):
                row = table.add_row()
                for j, cell_text in enumerate(cells):
                    row.cells[j].text = cell_text
            i += 1

        # Code block (```)
        elif stripped.startswith("```"):
            code_lang = stripped[3:].strip()
            # Collect code lines
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```

            # Add code paragraph
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)

        # Bullet list
        elif stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            # Handle bold inline (**text**)
            process_inline_formatting(p, text)
            i += 1

        # Numbered list
        elif re.match(r"^\d+\.\s", stripped):
            # Extract number prefix
            match = re.match(r"^(\d+\.)\s(.*)", stripped)
            if match:
                p = doc.add_paragraph(style="List Number")
                p.paragraph_format.left_indent = Inches(0.25)
                p.clear()
                p.add_run(f"{match.group(1)} ")
                run = p.add_run(match.group(2))
                i += 1

        # Empty line
        elif not stripped:
            i += 1

        # Regular paragraph
        else:
            # Skip table of contents lines (they have [text](#anchor))
            if stripped.startswith("[") and "](#" in stripped:
                i += 1
                continue

            p = doc.add_paragraph()
            process_inline_formatting(p, stripped)
            i += 1


def process_inline_formatting(para, text):
    """Handle **bold** and *italic* inline formatting."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = para.add_run(part[1:-1])
            run.italic = True
        else:
            para.add_run(part)


def main():
    print(f"Reading {INPUT_FILE}...")
    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        lines = f.readlines()

    print("Creating DOCX document...")
    doc = Document()

    # Document title
    title = doc.add_heading("BÁO CÁO THỰC HÀNH", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph("GIẢI THUẬT XỬ LÝ SONG SONG VÀ PHÂN BỔ")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.bold = True
        run.font.size = Pt(14)

    add_horizontal_line(doc)

    # Meta info
    doc.add_paragraph("Môn: Giải thuật xử lý song song và phân bổ")
    doc.add_paragraph("GVHD: Lê Minh Khánh Hội, Lê Phạm Hoàng Trung")
    doc.add_paragraph("Ngày: 31 tháng 3 năm 2026")

    add_horizontal_line(doc)

    # Process content
    process_markdown(doc, lines)

    # Save
    print(f"Saving to {OUTPUT_FILE}...")
    doc.save(OUTPUT_FILE)
    print(f"Done! Report saved to {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
